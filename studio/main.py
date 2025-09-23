import os
import uuid
from pathlib import Path
from typing import Dict, List
from starlette.requests import Request
from fastapi import FastAPI, Request, Form, Depends, Query
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse
from sqlmodel import Session, select
from starlette.middleware.sessions import SessionMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest
from starlette.middleware import Middleware
from docx import Document
import uvicorn
import logging

# Chat & model handling
from chat_client.generic_model_config import MODEL_REGISTRY, get_model_handler

# App components
from studio.routes import auth_routes, config_routes, chat_routes, newsdata_routes, project_and_settings_routes
from studio.services.db import init_db, get_session
from studio.dependency import get_current_user, require_roles
from studio.services.config_manager import load_user_config_db, save_user_config_db
from studio.services.db import get_user_sessions
# from studio.settings import DEFAULT_USER_CONFIG 

# Tool routes
from studio.auth import router as auth_router
from tools.context_manager.context_management_routes import router as context_router
from tools.RAG.rag_routes import router as rag_router
from tools.MARKDOWN.md_routes import router as md_router
from tools.DOCGEN.docgen_routes import router as docgen_router


logging.basicConfig(
    level=logging.DEBUG,  # or INFO
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)
logger = logging.getLogger(__name__)

logger.info("Logger initiated")

class AuthContextMiddleware(BaseHTTPMiddleware):
    logger.info(f"{__name__}\t- [initiated]")
    async def dispatch(self, request: StarletteRequest, call_next):
        logger.info(f"{__name__}\t- [initiated]")
        session = request.session  # Now this is safe
        user_id = session.get("user_id")
        if user_id:
            request.state.user = {"id": user_id, "name": "Test User"}
            print("Session user_id:", user_id)
        else:
            request.state.user = None
        return await call_next(request)

middleware = [
    Middleware(SessionMiddleware, secret_key="Shivaa@2025"),
    Middleware(AuthContextMiddleware)
]

# Initiating FastAPI app interfae post middleware defination. 
app = FastAPI(middleware=middleware)
logger.info(f"AI Studio\t- [initiated]")

# app.add_middleware(SessionMiddleware, secret_key="Shivaa@2025")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=str(Path(__file__).parent / "templates"))
app.mount("/static", StaticFiles(directory="studio/static"), name="static")
templates.env.cache.clear()

init_db()

# Include routers
app.include_router(auth_routes.router, tags=["Authentication"])     # Need special alignment.
app.include_router(config_routes.router, prefix="/config", tags=["Configuration"])
app.include_router(auth_router, prefix="/auth", tags=["OAuth"])     # Need special alignment.
app.include_router(newsdata_routes.router, tags=["News"])
app.include_router(project_and_settings_routes.router, tags=["Settings"])
app.include_router(chat_routes.router, prefix="/chat", tags=["Chat"])

# Configuration update
app.include_router(context_router, prefix="/tools/context_manager")
app.include_router(rag_router, prefix="/tools/rag")
app.include_router(md_router, prefix="/tools/markdown")
app.include_router(docgen_router, prefix="/tools/docgen")
    
conversations_store: Dict[str, List[Dict[str, str]]] = {}  # {conversation_name: [ {"user":..., "ai":...}, ... ]}

@app.get("/favicon.ico")
async def favicon():
    logger.info(f"favicon\t- [initiated]")
    return FileResponse("studio/static/favicon.ico")

@app.post("/generate-doc")
async def generate_doc(request: Request, content: str = Form(...)):
    logger.info(f"{__name__}\t- [initiated]")
    doc = Document()
    doc.add_heading("AI Response", level=1)
    doc.add_paragraph(content)

    filename = f"response_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(BASE_DIR, "generated", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)

    return FileResponse(path=filepath, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    logger.info(f"{__name__}\t- [initiated]")
    user = request.state.user
    if user:
        logger.info(f"\n\nuser=>\n{user}\nid=>{user['id']}\n\n~~~~")
        sessions = get_user_sessions(user['id'])  # fetch from DB
        logger.info(f"\n\sessions=>\n{sessions}\n\n~~~~")
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "models": MODEL_REGISTRY,
                "conversations": sessions,
            }
        )
    else:
        return RedirectResponse(url="/chat/public/")

@app.get("/me/config")
def get_config(
    user = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    logger.info(f"{__name__}\t- [initiated]")
    return load_user_config_db(session, user["id"], default={})

@app.put("/me/config")
def update_config(
    body: dict,
    user = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    logger.info(f"{__name__}\t- [initiated]")
    save_user_config_db(session, user["id"], body)
    return {"ok": True}

@app.get("/me")
def me(user = Depends(get_current_user)):
    logger.info(f"{__name__}\t- [initiated]")
    return {"id": user.id, "email": user.email, "role": user.role}

@app.get("/admin/metrics")
def metrics(user = Depends(require_roles("admin"))):
    logger.info(f"{__name__}\t- [initiated]")
    return {"secret": "..." }


if __name__ == "__main__":
    logger.info(f"{__name__}\t- [initiated]")
    uvicorn.run("studio.main:app", host="127.0.0.1", port=8030, reload=True)
